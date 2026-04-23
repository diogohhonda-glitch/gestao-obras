"""
Proposta Técnica/Comercial 344-2026 – Empreendimento SORANO
Imper Soluções Ltda.
Cores: Laranja #E8632A + Carvão #3D3D3D (conforme logo)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import os

# ── Cores da marca ─────────────────────────────────────────────────────────────
LARANJA     = RGBColor(0xE8, 0x63, 0x2A)   # laranja Imper
CARVAO      = RGBColor(0x3D, 0x3D, 0x3D)   # carvão/preto IMPER
CARVAO_CLR  = RGBColor(0x5A, 0x5A, 0x5A)   # carvão médio para corpo
BRANCO      = RGBColor(0xFF, 0xFF, 0xFF)
LARANJA_CLR = RGBColor(0xFF, 0xE4, 0xD0)   # laranja claro (fundo alternado)
CINZA_BG    = RGBColor(0xF5, 0xF5, 0xF5)   # cinza fundo
PRETO       = RGBColor(0x00, 0x00, 0x00)

# hex para shading XML
H_LARANJA     = 'E8632A'
H_CARVAO      = '3D3D3D'
H_LARANJA_CLR = 'FFEEE5'
H_CINZA_BG    = 'F5F5F5'
H_BRANCO      = 'FFFFFF'
H_LARANJA_MED = 'F28850'  # laranja médio para subheaders

# Logo path
LOGO_PATH = 'c:/Windows/Temp/Diogo/Sorano/extracted/word/media/image1.png'

doc = Document()

# ── Página ────────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Cm(21.0)
section.page_height   = Cm(29.7)
section.left_margin   = Cm(2.0)
section.right_margin  = Cm(2.0)
section.top_margin    = Cm(1.8)
section.bottom_margin = Cm(1.8)
section.header_distance = Cm(1.0)
section.footer_distance = Cm(1.0)

# ── Helpers ────────────────────────────────────────────────────────────────────
def sf(run, name='Calibri', sz=10, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(sz)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove old shd
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def no_border_cell(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        e.set(qn('w:color'), 'auto')
        tcB.append(e)
    tcPr.append(tcB)

def slim_border_cell(cell, color=H_LARANJA):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '4')
        e.set(qn('w:color'), color)
        tcB.append(e)
    tcPr.append(tcB)

def tc(cell, text='', bold=False, sz=9, color=CARVAO_CLR, bg=None,
       align=WD_ALIGN_PARAGRAPH.LEFT, valign=WD_ALIGN_VERTICAL.CENTER, italic=False):
    cell.vertical_alignment = valign
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    sf(run, sz=sz, bold=bold, color=color, italic=italic)
    if bg:
        shade_cell(cell, bg)

def divider(doc, color=H_LARANJA):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '8')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def hdr_para(doc, text, sz=13, color=CARVAO, space_before=12, space_after=6,
             align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    sf(run, sz=sz, bold=True, color=color, italic=italic)
    return p

def body(doc, text, sz=10, color=CARVAO_CLR, sb=2, sa=4,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    run = p.add_run(text)
    sf(run, sz=sz, bold=bold, color=color, italic=italic)
    return p

def bullet_item(doc, text, sz=10, color=CARVAO_CLR):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    sf(run, sz=sz, color=color)
    return p

def spacer(doc, h=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(h)
    p.paragraph_format.space_after  = Pt(0)

# ══════════════════════════════════════════════════════════════════════════════
# CABEÇALHO (Header) com logo + faixa de contato
# ══════════════════════════════════════════════════════════════════════════════
def build_header(doc, logo_path):
    hdr = doc.sections[0].header
    # Clear default paragraph
    for p in hdr.paragraphs:
        p.clear()

    # Use a 1-row, 2-col table: col0=logo, col1=contatos
    t = hdr.add_table(1, 2, hdr.paragraphs[0]._p)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.columns[0].width = Cm(6.0)
    t.columns[1].width = Cm(11.0)

    # Remove ALL borders
    for row in t.rows:
        for cell in row.cells:
            no_border_cell(cell)

    # Col 0 - Logo
    cell_logo = t.rows[0].cells[0]
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after  = Pt(0)
    run_logo = p_logo.add_run()
    run_logo.add_picture(logo_path, width=Cm(5.2))

    # Col 1 - Contatos (faixa laranja)
    cell_info = t.rows[0].cells[1]
    cell_info.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    shade_cell(cell_info, H_LARANJA)
    slim_border_cell(cell_info, H_LARANJA)

    contact_lines = [
        ('Imper Soluções Ltda.', True, 10),
        ('CNPJ: 31.527.551/0001-31', False, 8),
        ('comercial01@impersolucoes.com.br  |  www.impersolucoes.com.br', False, 8),
        ('@impersolucoes  |  Diogo Honda', False, 8),
        ('07/04/2026  –  Proposta Nº 344-2026', False, 8),
    ]
    for i, (line, bold, sz) in enumerate(contact_lines):
        if i == 0:
            p = cell_info.paragraphs[0]
        else:
            p = cell_info.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.3)
        run = p.add_run(line)
        sf(run, sz=sz, bold=bold, color=BRANCO)

build_header(doc, LOGO_PATH)

# ══════════════════════════════════════════════════════════════════════════════
# RODAPÉ (Footer)
# ══════════════════════════════════════════════════════════════════════════════
def build_footer(doc):
    ftr = doc.sections[0].footer
    for p in ftr.paragraphs:
        p.clear()
    p = ftr.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)

    # Linha divisória laranja
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), H_LARANJA)
    pBdr.append(top)
    pPr.append(pBdr)

    r1 = p.add_run('Imper Soluções Ltda.  |  ')
    sf(r1, sz=8, bold=True, color=CARVAO)
    r2 = p.add_run('CNPJ: 31.527.551/0001-31  |  comercial01@impersolucoes.com.br  |  www.impersolucoes.com.br')
    sf(r2, sz=8, color=CARVAO_CLR)

    # Page number
    p2 = ftr.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run('Página ')
    sf(r, sz=8, color=CARVAO_CLR)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2 = p2.add_run()
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)
    sf(run2, sz=8, color=CARVAO_CLR)

build_footer(doc)

# ══════════════════════════════════════════════════════════════════════════════
# CAPA / TÍTULO
# ══════════════════════════════════════════════════════════════════════════════

# Faixa carvão com título
p_capa = doc.add_paragraph()
p_capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_capa.paragraph_format.space_before = Pt(10)
p_capa.paragraph_format.space_after  = Pt(0)
# Fundo carvão via shading via table trick — use inline run styling
run = p_capa.add_run('PROPOSTA TÉCNICA E COMERCIAL')
sf(run, sz=20, bold=True, color=CARVAO)

p_num = doc.add_paragraph()
p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_num.paragraph_format.space_before = Pt(2)
p_num.paragraph_format.space_after  = Pt(2)
run = p_num.add_run('Nº 344-2026')
sf(run, sz=14, bold=True, color=LARANJA)

# Faixa laranja - nome do empreendimento
t_capa = doc.add_table(1, 1)
t_capa.alignment = WD_TABLE_ALIGNMENT.CENTER
t_capa.columns[0].width = Cm(17.0)
cell_emp = t_capa.rows[0].cells[0]
shade_cell(cell_emp, H_CARVAO)
no_border_cell(cell_emp)
p_emp = cell_emp.paragraphs[0]
p_emp.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_emp.paragraph_format.space_before = Pt(8)
p_emp.paragraph_format.space_after  = Pt(8)
r1 = p_emp.add_run('Empreendimento  ')
sf(r1, sz=14, bold=False, color=BRANCO)
r2 = p_emp.add_run('SORANO')
sf(r2, sz=16, bold=True, color=LARANJA)
r3 = p_emp.add_run('  |  Curitiba – PR')
sf(r3, sz=12, bold=False, color=BRANCO)

spacer(doc, 10)
divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# BLOCOS CONTRATANTE / CONTRATADA
# ══════════════════════════════════════════════════════════════════════════════
# Tabela 2 colunas lado a lado
t_partes = doc.add_table(1, 2)
t_partes.alignment = WD_TABLE_ALIGNMENT.CENTER
t_partes.columns[0].width = Cm(8.3)
t_partes.columns[1].width = Cm(8.7)

cell_cte = t_partes.rows[0].cells[0]
cell_ctd = t_partes.rows[0].cells[1]
no_border_cell(cell_cte)
no_border_cell(cell_ctd)

def info_block(cell, titulo, items, bg=H_CARVAO):
    # Header row dentro da célula via paragraphs
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f'  {titulo}')
    sf(run, sz=11, bold=True, color=LARANJA)

    shade_cell(cell, H_BRANCO)
    # Draw orange top bar
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), H_LARANJA)
    pBdr.append(top)
    pPr.append(pBdr)

    for label, value in items:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(1)
        p2.paragraph_format.left_indent  = Cm(0.3)
        r1 = p2.add_run(f'{label} ')
        sf(r1, sz=9, bold=True, color=CARVAO)
        r2 = p2.add_run(value)
        sf(r2, sz=9, color=CARVAO_CLR)

# Contratante
info_block(cell_cte, 'CONTRATANTE', [
    ('Empreendimento:', 'SORANO'),
    ('Endereço:', 'Av. República Argentina, 4483, Curitiba – PR'),
    ('Contato:', ''),
    ('Telefone:', ''),
    ('E-mail:', ''),
    ('Data Emissão:', '07/04/2026'),
    ('Validade:', '30 dias corridos'),
])

# Contratada
info_block(cell_ctd, 'CONTRATADA', [
    ('Empresa:', 'Imper Soluções Ltda.'),
    ('CNPJ:', '31.527.551/0001-31'),
    ('Responsável:', 'Diogo Honda'),
    ('Telefone:', ''),
    ('E-mail:', 'comercial01@impersolucoes.com.br'),
    ('Site:', 'www.impersolucoes.com.br'),
    ('Instagram:', '@impersolucoes'),
])

spacer(doc, 10)
divider(doc)

# ── Apresentação ──────────────────────────────────────────────────────────────
body(doc,
    'Prezados Senhores,\n\n'
    'Agradecemos a oportunidade de apresentar nossa Proposta Técnica e Comercial para a '
    'execução dos serviços de impermeabilização do Empreendimento SORANO. A Imper Soluções '
    'possui experiência consolidada na execução de sistemas de impermeabilização de alto desempenho '
    'em obras de grande porte, operando com equipe especializada, materiais certificados e em '
    'conformidade plena com as normas técnicas ABNT vigentes.\n\n'
    'Esta proposta foi elaborada com base no Memorial Descritivo de Impermeabilização (Imperproject – '
    'Revisão V2, 09/04/2025) e no Procedimento de Execução ITS.008, contemplando todos os sistemas '
    'especificados em projeto, divididos em quatro etapas conforme cronograma da obra.',
    sb=6, sa=10)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SEÇÃO: helper para títulos de seção
# ══════════════════════════════════════════════════════════════════════════════
def section_title(doc, text, sb=10):
    """Faixa carvão com texto laranja – título principal de seção"""
    t = doc.add_table(1, 1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.columns[0].width = Cm(17.0)
    cell = t.rows[0].cells[0]
    shade_cell(cell, H_CARVAO)
    no_border_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.4)
    run = p.add_run(text)
    sf(run, sz=11, bold=True, color=LARANJA)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def sub_title(doc, text, sb=8, sa=4):
    """Linha laranja como subtítulo"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), H_LARANJA)
    pBdr.append(left)
    pPr.append(pBdr)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    sf(run, sz=11, bold=True, color=CARVAO)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# 1. OBJETO / ESCOPO
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '1.  OBJETO E ESCOPO DO SERVIÇO')

body(doc, 'O objeto desta proposta é a execução completa dos serviços de '
     'impermeabilização do Empreendimento SORANO, compreendendo:', sa=4)

for item in [
    'Preparação e regularização de superfícies horizontais e verticais conforme especificação de projeto;',
    'Impermeabilização de fossos de elevadores com argamassa polimérica – Viaplus 1000 (4,00 kg/m²);',
    'Impermeabilização de depósitos de lixo e áreas úmidas internas com argamassa polimérica (4,00 kg/m²);',
    'Impermeabilização de sacadas, varandas e gardens com membrana acrílica bicomponente '
     'com cimento e fibras sintéticas – Viaplus 7000 (3,00 e 4,50 kg/m²);',
    'Pintura antiraiz com emulsão asfáltica – Viabit Acqua (0,80 L/m²) em jardineiras e gardens;',
    'Impermeabilização de lajes de cobertura com manta asfáltica aderida AA 4mm, Tipo III, '
     'Classe B – Torodin, com primer Ecoprimer e asfalto oxidado modificado NBR II/III;',
    'Impermeabilização de reservatórios com membrana à base de poliuretano – Vitpoli Eco (2,00 kg/m²) '
     'e/ou epóxi poliamida – Viapoxi Coat (1,00 kg/m²);',
    'Impermeabilização complementar com membrana acrílica sem cimento – Vialastic (1,00 L/m²);',
    'Fornecimento e aplicação de tela de reforço de poliéster com PVC – B-Tuk nos pontos críticos;',
    'Fornecimento e aplicação de camada separadora em filme de polietileno 24 micras;',
    'Execução de testes de estanqueidade – lâmina d\'água mínima de 1 cm por 72 horas (NBR 9574);',
    'Proteção mecânica horizontal e vertical conforme especificação do projeto.',
]:
    bullet_item(doc, item)

body(doc,
    'Os quantitativos foram levantados com base no Memorial Descritivo (BIM – Revisão V2, Imperproject) '
    'e estão sujeitos a conferência presencial antes do início de cada etapa.',
    sz=9, italic=True, sb=4, sa=8)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 2. CRONOGRAMA
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '2.  CRONOGRAMA DE EXECUÇÃO POR ETAPAS')

body(doc,
    'Os serviços serão executados em quatro etapas, conforme as datas de necessidade '
    'definidas pela Contratante:', sa=6)

# Tabela cronograma
t_cr = doc.add_table(6, 5)
t_cr.style = 'Table Grid'
t_cr.alignment = WD_TABLE_ALIGNMENT.CENTER

widths_cr = [Cm(2.0), Cm(3.8), Cm(5.8), Cm(2.8), Cm(2.6)]
for ci, w in enumerate(widths_cr):
    for row in t_cr.rows:
        row.cells[ci].width = w

# Header
hcols = ['Etapa', 'Descrição', 'Áreas Envolvidas', 'Início Previsto', 'Necessidade']
for ci, h in enumerate(hcols):
    tc(t_cr.rows[0].cells[ci], h, bold=True, sz=9, color=BRANCO, bg=H_CARVAO,
       align=WD_ALIGN_PARAGRAPH.CENTER)

etapas = [
    ('1ª Etapa', 'Reservatórios\nTérreo',
     'Fossos de elevadores (4 fossos), depósitos de lixo, reservatórios e calhas',
     '08/06/2026', '13/07/2026'),
    ('2ª Etapa', 'Gardens',
     'Gardens, jardineiras e sacadas – Blocos 1 e 2 (2º ao 18º pavimento)',
     '20/07/2026', '18/08/2026'),
    ('3ª Etapa', 'Área\nExterna',
     'Áreas externas, varandas abertas, lajes intermediárias e calhas expostas',
     '17/08/2026', '16/09/2026'),
    ('4ª Etapa', 'Cobertura e\nReservatórios',
     'Lajes de cobertura (manta asfáltica AA 4mm) e reservatórios superiores',
     '01/12/2026', '11/01/2027'),
]

for ri, (et, desc, areas, inicio, nec) in enumerate(etapas):
    row = t_cr.rows[ri + 1]
    bg = H_LARANJA_CLR if ri % 2 == 0 else H_BRANCO
    tc(row.cells[0], et, bold=True, sz=9, color=LARANJA, bg=H_CARVAO, align=WD_ALIGN_PARAGRAPH.CENTER)
    tc(row.cells[1], desc, bold=True, sz=9, color=CARVAO, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    tc(row.cells[2], areas, sz=9, color=CARVAO_CLR, bg=bg)
    tc(row.cells[3], inicio, sz=9, color=CARVAO_CLR, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    tc(row.cells[4], nec, bold=True, sz=9, color=LARANJA, bg=H_CARVAO, align=WD_ALIGN_PARAGRAPH.CENTER)

# Total row
row_tot = t_cr.rows[5]
row_tot.cells[0].merge(row_tot.cells[3])
tc(row_tot.cells[0], 'Duração total da obra (todas as etapas):',
   bold=True, sz=9, color=BRANCO, bg=H_CARVAO, align=WD_ALIGN_PARAGRAPH.RIGHT)
tc(row_tot.cells[4], 'Jun/2026 → Jan/2027',
   bold=True, sz=9, color=LARANJA, bg=H_CARVAO, align=WD_ALIGN_PARAGRAPH.CENTER)

spacer(doc, 4)
body(doc,
    '⚠  Datas de início previsto dependem da liberação de frentes de trabalho com antecedência mínima '
    'de 10 dias úteis. Condições climáticas adversas ou alterações de projeto serão comunicadas '
    'formalmente à Contratante e não configurarão atraso por parte da Contratada.',
    sz=9, italic=True, sb=4, sa=8)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 3. DESCRIÇÃO TÉCNICA
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '3.  DESCRIÇÃO TÉCNICA DOS SERVIÇOS POR ETAPA')

# ── ETAPA 1 ──
sub_title(doc, '3.1  1ª Etapa – Reservatórios Térreo  |  Necessidade: 13/07/2026')
body(doc, 'Áreas: Fossos de elevadores, depósitos de lixo e reservatórios do pavimento térreo.',
     bold=True, sa=2)
body(doc, 'Sistema: Argamassa Polimérica Industrializada – Viaplus 1000 Viapol  |  Consumo: 4,00 kg/m²', sa=3)
for b in [
    'Regularização e limpeza da superfície: remoção de poeira, nata, óleo, graxa e materiais soltos; '
     'recomposição de fissuras com graute ou argamassa 1:3;',
    'Umedecimento da superfície até estado "saturado seco"; mistura mecânica dos componentes A+B por 3–5 min;',
    'Aplicação em demãos cruzadas com trincha/brocha – 1,00 kg/m² por demão, intervalo de 4 a 6 horas, '
     'totalizando 4,00 kg/m²;',
    'Reforço com tela de poliéster B-Tuk em cantos, ralos e junções piso/parede;',
    'Teste de estanqueidade: lâmina d\'água ≥ 1 cm por 72 horas;',
    'Proteção mecânica com chapisco e argamassa após aprovação do teste.',
]:
    bullet_item(doc, b)

# ── ETAPA 2 ──
sub_title(doc, '3.2  2ª Etapa – Gardens  |  Necessidade: 18/08/2026')
body(doc, 'Áreas: Gardens, jardineiras, sacadas e varandas – Blocos 1 e 2 (2º ao 18º pavimento).',
     bold=True, sa=2)
body(doc, 'Sistemas:\n'
     '① Membrana Acrílica c/ Cimento e Fibras – Viaplus 7000 Viapol  |  3,00 kg/m² (padrão) / 4,50 kg/m² (reforçado)\n'
     '② Pintura Antiraiz – Viabit Acqua Viapol  |  0,80 L/m²', sa=3)
for b in [
    'Preparação: raspagem, remoção de resíduos, recomposição de falhas; verificação de caimentos '
     '(mín. 0,5% interno, 1,0% externo, 2,0% para membrana acrílica);',
    'Rebaixo de 1 cm nos ralos (área 40×40 cm); arredondamento de cantos e arestas (Ø mín. 5 cm);',
    'Aplicação da membrana em demãos cruzadas com brocha – paredes primeiro (mín. 1,50 m no box), depois piso;',
    'Tela de poliéster no rodapé (10 cm parede + 10 cm piso), cantos internos e ralos (30×30 cm);',
    '2ª demão após cura mínima de 3 horas (ou conforme fabricante);',
    'Pintura antiraiz Viabit Acqua com rolo de lã sobre proteção mecânica nas jardineiras;',
    'Teste de estanqueidade no box e sacada: lâmina d\'água ≥ 1 cm por 72 horas.',
]:
    bullet_item(doc, b)

# ── ETAPA 3 ──
sub_title(doc, '3.3  3ª Etapa – Área Externa  |  Necessidade: 16/09/2026')
body(doc, 'Áreas: Áreas externas, lajes intermediárias expostas, calhas e elementos drenantes.',
     bold=True, sa=2)
body(doc, 'Sistemas:\n'
     '① Membrana Acrílica sem Cimento – Vialastic Viapol  |  1,00 L/m²\n'
     '② Membrana Acrílica c/ Cimento e Fibras – Viaplus 7000 (áreas de maior solicitação)', sa=3)
for b in [
    'Limpeza e preparação conforme itens 2.1 e 2.2 do Memorial Descritivo;',
    'Vialastic: 1ª demão diluída 1:1 sobre superfície seca; demãos seguintes puras (300 ml/m²), '
     'em sentidos cruzados; intervalo de 8 horas entre demãos;',
    'Reforço com tela B-Tuk em pontos críticos, juntas e mudanças de plano;',
    'Camada separadora em filme de polietileno 24 micras entre impermeabilização e proteção mecânica;',
    'Preenchimento de juntas perimetrais (2 cm) com argamassa betuminosa 1:8:3;',
    'Proteção mecânica primária, convencional ou reforçada com fibras conforme especificação de projeto;',
    'Teste de estanqueidade com lâmina d\'água (72 horas) nas áreas de maior exposição.',
]:
    bullet_item(doc, b)

# ── ETAPA 4 ──
sub_title(doc, '3.4  4ª Etapa – Cobertura e Reservatórios  |  Necessidade: 11/01/2027')
body(doc, 'Áreas: Lajes de cobertura (manta asfáltica) e reservatórios superiores.',
     bold=True, sa=2)
body(doc, 'Sistemas:\n'
     '① Manta Asfáltica AA 4mm – Torodin Tipo III Cl.B + Ecoprimer + Asfalto NBR II/III – Cobertura\n'
     '② Membrana Poliuretano – Vitpoli Eco Viapol  |  2,00 kg/m² – Reservatórios\n'
     '③ Membrana Epóxi Poliamida – Viapoxi Coat Viapol  |  1,00 kg/m² – Reservatórios (onde especificado)', sa=3)

body(doc, 'Manta Asfáltica – Cobertura:', bold=True, sz=9, sb=4, sa=2)
for b in [
    'Preparação horizontal: arredondamento de cantos/arestas (Ø mín. 5 cm);',
    'Primer Ecoprimer: 0,40 L/m² com rolo de lã; secagem mínima de 6 horas;',
    'Posicionamento dos rolos no sentido contrário ao fluxo de água, da cota mais baixa para a mais alta;',
    'Aquecimento do asfalto oxidado em caldeira elétrica a 180–200°C;',
    'Colagem com rolete metálico (centro→bordas); sobreposição mínima de 10 cm com banho de asfalto;',
    'Rodapé colado com altura mínima de 30 cm; sobreposição mínima de 10 cm sobre manta de piso;',
    'Camada separadora (filme PE 24 micras) + proteção mecânica reforçada (microconcreto fck 30 MPa, 5 cm);',
    'Teste de estanqueidade 72 horas após cura completa.',
]:
    bullet_item(doc, b)

body(doc, 'Reservatórios – Poliuretano / Epóxi:', bold=True, sz=9, sb=4, sa=2)
for b in [
    'Regularização da superfície: trincas, fissuras e ferros expostos tratados; lavagem completa;',
    'Primer à base de poliuretano vegetal: 0,30 kg/m²; aguardar 3 horas;',
    'Tela de poliéster fixada nos rodapés, junções e juntas frias com película ainda adesiva;',
    'Aplicação em demãos cruzadas (intervalo de 6 horas) até 2,00 kg/m²;',
    'Areia seca peneirada espalhada antes da secagem final para aderência da proteção;',
    'Cura mínima de 7 dias antes de qualquer solicitação mecânica;',
    'Ralo impermeabilizado com tela em "margarida" (3 cm de entrada);',
    'Proteção mecânica com chapisco + argamassa fina; contrapiso quando especificado.',
]:
    bullet_item(doc, b)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 4. VALORES
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '4.  VALORES E CONDIÇÕES COMERCIAIS')

body(doc,
    'Os valores abaixo contemplam fornecimento de todos os materiais especificados em projeto, '
    'mão de obra especializada, testes de estanqueidade, proteção mecânica e limpeza final de cada etapa.',
    sa=8)

t_val = doc.add_table(7, 4)
t_val.style = 'Table Grid'
t_val.alignment = WD_TABLE_ALIGNMENT.CENTER
val_widths = [Cm(2.0), Cm(6.0), Cm(4.0), Cm(5.0)]
for ci, w in enumerate(val_widths):
    for row in t_val.rows:
        row.cells[ci].width = w

val_hdrs = ['Etapa', 'Descrição', 'Área Aprox. (m²)', 'Valor (R$)']
for ci, h in enumerate(val_hdrs):
    tc(t_val.rows[0].cells[ci], h, bold=True, sz=10, color=BRANCO, bg=H_CARVAO,
       align=WD_ALIGN_PARAGRAPH.CENTER)

val_rows = [
    ('1ª', 'Reservatórios Térreo', '', ''),
    ('2ª', 'Gardens', '', ''),
    ('3ª', 'Área Externa', '', ''),
    ('4ª', 'Cobertura e Reservatórios', '', ''),
    ('', 'TOTAL GERAL', '', ''),
]
bgs_val  = [H_LARANJA_CLR, H_BRANCO, H_LARANJA_CLR, H_BRANCO, H_CARVAO]
clrs_val = [CARVAO, CARVAO_CLR, CARVAO, CARVAO_CLR, BRANCO]
for ri, ((et, desc, area, val), bg, col) in enumerate(zip(val_rows, bgs_val, clrs_val)):
    row = t_val.rows[ri + 1]
    is_tot = ri == 4
    tc(row.cells[0], et, bold=is_tot, sz=10, color=LARANJA if is_tot else CARVAO,
       bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    tc(row.cells[1], desc, bold=is_tot, sz=10, color=col, bg=bg)
    tc(row.cells[2], area, sz=10, color=col, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    tc(row.cells[3], val, bold=is_tot, sz=11, color=LARANJA if is_tot else col,
       bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)

# Linha extra: BDI info
row_bdi = t_val.rows[6]
row_bdi.cells[0].merge(row_bdi.cells[3])
tc(row_bdi.cells[0],
   '* Valores sujeitos a medição presencial de cada etapa. Não incluem regularização estrutural '
   'nem retrabalhos por falhas de terceiros.',
   sz=8, color=CARVAO_CLR, bg=H_CINZA_BG, italic=True)

spacer(doc, 6)

# Forma de pagamento
sub_title(doc, '4.1  Forma de Pagamento', sb=6)
body(doc, 'Proposta a definir em negociação. Modelo sugerido:', sa=3)
for b in [
    '30% de entrada na assinatura do contrato de cada etapa;',
    '40% na conclusão de 50% da área impermeabilizada da etapa;',
    '30% após aprovação do teste de estanqueidade e entrega formal da etapa;',
    'Faturamento via Nota Fiscal com dados completos da Contratante.',
]:
    bullet_item(doc, b)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 5. EQUIPE
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '5.  EQUIPE TÉCNICA E MOBILIZAÇÃO')

body(doc,
    'A Imper Soluções mobilizará para cada etapa equipe devidamente qualificada, '
    'com documentação e treinamentos em vigência:', sa=6)

t_eq = doc.add_table(4, 3)
t_eq.style = 'Table Grid'
t_eq.alignment = WD_TABLE_ALIGNMENT.CENTER
t_eq.columns[0].width = Cm(4.5)
t_eq.columns[1].width = Cm(3.5)
t_eq.columns[2].width = Cm(9.0)

for ci, h in enumerate(['Função', 'Quantidade', 'Qualificação / Treinamentos Obrigatórios']):
    tc(t_eq.rows[0].cells[ci], h, bold=True, sz=9, color=BRANCO, bg=H_CARVAO,
       align=WD_ALIGN_PARAGRAPH.CENTER)

equipe = [
    ('Supervisor Técnico', '1 por etapa', 'NR-18, NR-35, NR-6, ASO válido, PCMSO, registro CREA ativo'),
    ('Impermeabilizador Oficial', '3 a 4 por etapa', 'NR-18, NR-35 (trabalho em altura), ASO válido'),
    ('Auxiliar de Obras', '1 a 2 por etapa', 'NR-18, ASO válido'),
]
for ri, (f, q, qual) in enumerate(equipe):
    row = t_eq.rows[ri + 1]
    bg = H_LARANJA_CLR if ri % 2 == 0 else H_BRANCO
    tc(row.cells[0], f, bold=True, sz=9, color=CARVAO, bg=bg)
    tc(row.cells[1], q, sz=9, color=CARVAO_CLR, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    tc(row.cells[2], qual, sz=9, color=CARVAO_CLR, bg=bg)

spacer(doc, 4)
body(doc,
    'Todos os colaboradores utilizarão EPIs completos: capacete, óculos, luvas, bota de segurança, '
    'cinto paraquedista e trava-queda para trabalho em altura, conforme NR-6 e NR-35.',
    sz=9, sa=8)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 6. PRAZOS
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '6.  PRAZOS DE EXECUÇÃO')

t_pz = doc.add_table(6, 4)
t_pz.style = 'Table Grid'
t_pz.alignment = WD_TABLE_ALIGNMENT.CENTER
t_pz.columns[0].width = Cm(5.0)
t_pz.columns[1].width = Cm(3.5)
t_pz.columns[2].width = Cm(3.5)
t_pz.columns[3].width = Cm(5.0)

for ci, h in enumerate(['Etapa', 'Prazo de Execução', 'Teste Estanqueidade', 'Observação']):
    tc(t_pz.rows[0].cells[ci], h, bold=True, sz=9, color=BRANCO, bg=H_CARVAO,
       align=WD_ALIGN_PARAGRAPH.CENTER)

prazos = [
    ('1ª Etapa – Reservatórios Térreo', '≈ 25 dias corridos', '72 h (incluso)', 'Cura mín. 7 dias para PU'),
    ('2ª Etapa – Gardens', '≈ 25 dias corridos', '72 h/bloco (incluso)', 'Execução por blocos conforme liberação'),
    ('3ª Etapa – Área Externa', '≈ 22 dias corridos', '72 h (incluso)', 'Sujeito a condições climáticas'),
    ('4ª Etapa – Cobertura e Reservatórios', '≈ 35 dias corridos', '72 h (incluso)', 'Cura manta 7 dias antes da proteção'),
    ('PRAZO TOTAL', 'Jun/2026 → Jan/2027', '—', 'Com intervalos entre etapas'),
]
for ri, (et, pr, tst, obs) in enumerate(prazos):
    row = t_pz.rows[ri + 1]
    is_tot = ri == 4
    bg = H_CARVAO if is_tot else (H_LARANJA_CLR if ri % 2 == 0 else H_BRANCO)
    col = LARANJA if is_tot else CARVAO
    col2 = BRANCO if is_tot else CARVAO_CLR
    tc(row.cells[0], et, bold=is_tot, sz=9, color=col if is_tot else CARVAO, bg=bg)
    tc(row.cells[1], pr, bold=is_tot, sz=9, color=col, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    tc(row.cells[2], tst, sz=9, color=col2, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    tc(row.cells[3], obs, sz=8, color=col2, bg=bg)

spacer(doc, 4)
body(doc,
    'Prazos consideram jornada de segunda a sábado (8h/dia) em condições climáticas favoráveis. '
    'Períodos de chuva ou indisponibilidade de frentes de trabalho pela Contratante serão '
    'descontados do cômputo de dias.',
    sz=9, italic=True, sa=8)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 7. RESPONSABILIDADES
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '7.  RESPONSABILIDADES DAS PARTES')

sub_title(doc, '7.1  Responsabilidades da Contratante', sb=6)
for b in [
    'Fornecer ponto de energia elétrica próximo ao local do serviço;',
    'Fornecer ponto hídrico para lavagem das superfícies antes da aplicação;',
    'Disponibilizar local coberto e seguro para armazenamento de materiais e equipamentos;',
    'Disponibilizar sanitários, vestiários e área de alimentação para os colaboradores;',
    'Garantir acesso seguro às frentes de trabalho em todas as etapas;',
    'Fornecer andaimes, plataformas, balancins ou outros equipamentos de acesso certificados;',
    'Responsabilizar-se pelo descarte correto das embalagens e resíduos (Lei 12.305/2010 – PNRS);',
    'Comunicar com antecedência mínima de 10 dias úteis a liberação de cada frente de trabalho;',
    'Garantir que o substrato possua cura mínima de 28 dias (concreto e alvenaria).',
]:
    bullet_item(doc, b)

sub_title(doc, '7.2  Responsabilidades da Contratada', sb=8)
for b in [
    'Executar os serviços em conformidade com o Memorial Descritivo (Imperproject – Rev. V2) e ITS.008;',
    'Fornecer todos os materiais especificados com fichas técnicas e laudos de conformidade;',
    'Manter equipe com documentação e treinamentos de segurança em vigor (NR-6, NR-18, NR-35, ASO, PCMSO);',
    'Entregar cada etapa com laudo de estanqueidade aprovado (72 horas);',
    'Entregar o local limpo e organizado ao final de cada etapa;',
    'Comunicar imediatamente qualquer anomalia no substrato que comprometa o sistema;',
    'Emitir ART ao término dos serviços;',
    'Fornecer laudo fotográfico de execução de cada etapa.',
]:
    bullet_item(doc, b)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 8. GARANTIA
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '8.  PRAZO E CONDIÇÕES DE GARANTIA')

body(doc, 'A Imper Soluções concede garantia sobre os serviços executados conforme o quadro abaixo:', sa=6)

t_gar = doc.add_table(5, 3)
t_gar.style = 'Table Grid'
t_gar.alignment = WD_TABLE_ALIGNMENT.CENTER
t_gar.columns[0].width = Cm(7.5)
t_gar.columns[1].width = Cm(3.0)
t_gar.columns[2].width = Cm(6.5)

for ci, h in enumerate(['Sistema de Impermeabilização', 'Garantia', 'Condição']):
    tc(t_gar.rows[0].cells[ci], h, bold=True, sz=9, color=BRANCO, bg=H_CARVAO,
       align=WD_ALIGN_PARAGRAPH.CENTER)

garantias = [
    ('Argamassa Polimérica – Viaplus 1000', '5 anos', 'A partir da emissão da NF da etapa'),
    ('Membrana Acrílica c/ Cimento e Fibras – Viaplus 7000', '5 anos', 'A partir da emissão da NF da etapa'),
    ('Manta Asfáltica AA 4mm – Torodin', '10 anos', 'A partir da emissão da NF da etapa'),
    ('Membrana Poliuretano / Epóxi Poliamida', '5 anos', 'A partir da emissão da NF da etapa'),
]
for ri, (sis, prazo, cond) in enumerate(garantias):
    row = t_gar.rows[ri + 1]
    bg = H_LARANJA_CLR if ri % 2 == 0 else H_BRANCO
    tc(row.cells[0], sis, sz=9, color=CARVAO, bg=bg)
    tc(row.cells[1], prazo, bold=True, sz=11, color=LARANJA, bg=bg, align=WD_ALIGN_PARAGRAPH.CENTER)
    tc(row.cells[2], cond, sz=9, color=CARVAO_CLR, bg=bg)

spacer(doc, 6)
body(doc, 'A garantia é válida sob as seguintes condições:', sa=3)
for b in [
    'Manutenções preventivas realizadas conforme Plano de Manutenção do Memorial Descritivo (Imperproject);',
    'Nenhuma intervenção sobre a impermeabilização sem autorização técnica escrita da Imper Soluções;',
    'Intervenções de terceiros sem acompanhamento da Contratada causam perda imediata da garantia;',
    'A garantia não cobre danos por falhas estruturais, recalques, acidentes ou vandalismo;',
    'A garantia não cobre falhas decorrentes do descumprimento das responsabilidades da Contratante.',
]:
    bullet_item(doc, b)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 9. NORMAS
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '9.  NORMAS E REFERÊNCIAS TÉCNICAS ADOTADAS')

normas = [
    'ABNT NBR 9574:2008 – Execução de Impermeabilização;',
    'ABNT NBR 9575:2010 – Impermeabilização: Seleção e Projeto;',
    'ABNT NBR 15575:2013 – Edificações Habitacionais – Desempenho;',
    'ABNT NBR 11905:2015 – Argamassa Polimérica Industrializada para Impermeabilização;',
    'ABNT NBR 15885:2010 – Membrana de Polímero Acrílico com ou sem Cimento;',
    'ABNT NBR 13321:2023 – Membrana Acrílica para Impermeabilização;',
    'ABNT NBR 15487-1:2023 – Membrana de Poliuretano para Impermeabilização;',
    'ABNT NBR 16072:2012 – Argamassa Impermeável;',
    'ABNT NBR 9686:2006 – Solução Asfáltica para Imprimação;',
    'Guia de Diretrizes IBI – Instituto Brasileiro de Impermeabilização (2023);',
    'NR-6 – Equipamentos de Proteção Individual;  NR-18 – Condições de Trabalho na Construção;  NR-35 – Trabalho em Altura;',
    'Memorial Descritivo – Imperproject, Revisão V2, 09/04/2025;',
    'Procedimento de Execução ITS.008 – Revisão 09, 04/04/2025.',
]
for n in normas:
    bullet_item(doc, n, sz=9)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 10. INFORMAÇÕES COMPLEMENTARES
# ══════════════════════════════════════════════════════════════════════════════
section_title(doc, '10.  INFORMAÇÕES COMPLEMENTARES')

body(doc,
    'Esta proposta tem validade de 30 dias corridos a partir da data de emissão. Após esse prazo, '
    'valores e condições poderão ser revisados em função de variações de custo de materiais e mão de obra.',
    sa=4)
body(doc,
    'Serviços não previstos nesta proposta, identificados durante a execução, serão orçados '
    'separadamente e somente executados mediante aprovação prévia e formalização via aditivo contratual.',
    sa=4)
body(doc,
    'A Imper Soluções está à inteira disposição para esclarecimentos, reuniões técnicas e '
    'adequações necessárias ao presente escopo.',
    sa=16)

divider(doc)

# ══════════════════════════════════════════════════════════════════════════════
# ASSINATURAS
# ══════════════════════════════════════════════════════════════════════════════
body(doc, 'Atenciosamente,', sz=11, align=WD_ALIGN_PARAGRAPH.CENTER, sa=14)

t_ass = doc.add_table(3, 2)
t_ass.alignment = WD_TABLE_ALIGNMENT.CENTER
t_ass.columns[0].width = Cm(8.5)
t_ass.columns[1].width = Cm(8.5)

for row in t_ass.rows:
    for cell in row.cells:
        no_border_cell(cell)

tc(t_ass.rows[0].cells[0], '_________________________________',
   sz=10, color=CARVAO_CLR, align=WD_ALIGN_PARAGRAPH.CENTER)
tc(t_ass.rows[0].cells[1], '_________________________________',
   sz=10, color=CARVAO_CLR, align=WD_ALIGN_PARAGRAPH.CENTER)
tc(t_ass.rows[1].cells[0], 'Diogo Honda',
   bold=True, sz=11, color=LARANJA, align=WD_ALIGN_PARAGRAPH.CENTER)
tc(t_ass.rows[1].cells[1], 'Representante Legal – Contratante',
   bold=True, sz=10, color=CARVAO, align=WD_ALIGN_PARAGRAPH.CENTER)
tc(t_ass.rows[2].cells[0], 'Imper Soluções Ltda.  |  CNPJ: 31.527.551/0001-31',
   sz=9, color=CARVAO_CLR, align=WD_ALIGN_PARAGRAPH.CENTER)
tc(t_ass.rows[2].cells[1], 'SORANO – Curitiba – PR  |  Data: ___ / ___ / ______',
   sz=9, color=CARVAO_CLR, align=WD_ALIGN_PARAGRAPH.CENTER)

# ── Save ──────────────────────────────────────────────────────────────────────
out = 'c:/Windows/Temp/Diogo/Sorano/Proposta 344-2026 SORANO.docx'
doc.save(out)
print(f'Saved: {out}')
print(f'Size: {os.path.getsize(out)/1024:.0f} KB')
